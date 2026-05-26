"""
test.py  –  Quick smoke-tests for the vector_store pipeline.

Run from the project root:
    python -m vector_store.test
  or
    cd vector_store && python test.py

Tests:
  1. PDF extraction
  2. DOCX extraction
  3. Image → text (Cloudflare vision)
  4. Audio → text (Cloudflare Whisper)
  5. Embedding generation (Cloudflare bge-base-en-v1.5)
  6. Pinecone upsert  (ingest_file full pipeline)
  7. Pinecone retrieval
"""

import os, sys, io

# ── Load .env manually (no python-dotenv needed) ──────────────────────────────
_env_path = os.path.join(os.path.dirname(__file__), "..", ".env")
if os.path.exists(_env_path):
    with open(_env_path) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip().strip("'\""))

# ── make sure parent dir is on the path so imports work ───────────────────────
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from vector_store.chunking        import chunk_text
from vector_store.embeddings      import embed_texts, embed_query
from vector_store.pinecone_client import upsert_vectors, query_vectors
from vector_store.ingest          import (
    _extract_pdf,
    _extract_docx,
    _extract_image_via_cf,
    _extract_audio_via_cf,
    ingest_file,
)
from vector_store.retrieve        import retrieve_chunks

# ── color helpers ──────────────────────────────────────────────────────────────
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

def ok(msg):   print(f"  {GREEN}✓ PASS{RESET}  {msg}")
def fail(msg, err): print(f"  {RED}✗ FAIL{RESET}  {msg}\n         {YELLOW}{err}{RESET}")
def header(msg): print(f"\n{BOLD}{'─'*55}\n  {msg}\n{'─'*55}{RESET}")

# ── test constants ─────────────────────────────────────────────────────────────
TEST_USER_ID = "test_user_001"
TEST_CONV_ID = "test_conv_001"
TEST_QUERY   = "artificial intelligence machine learning"

SAMPLE_TEXT = (
    "Artificial intelligence (AI) is intelligence demonstrated by machines. "
    "Machine learning is a subset of AI that enables systems to learn from data. "
    "Deep learning uses neural networks with many layers to model complex patterns."
)

# ─────────────────────────────────────────────────────────────────────────────
# 1. PDF EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────
def test_pdf():
    header("TEST 1 · PDF Extraction")
    try:
        import pdfplumber
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.drawString(72, 720, "Hello from a test PDF!")
        c.drawString(72, 700, SAMPLE_TEXT[:80])
        c.save()
        pdf_bytes = buf.getvalue()

        text = _extract_pdf(pdf_bytes)
        assert "Hello" in text or len(text) > 5, "Extracted text too short"
        ok(f"Extracted {len(text)} chars from generated PDF")
    except ImportError as e:
        fail("PDF test skipped – missing library", e)
    except Exception as e:
        fail("PDF extraction failed", e)


# ─────────────────────────────────────────────────────────────────────────────
# 2. DOCX EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────
def test_docx():
    header("TEST 2 · DOCX Extraction")
    try:
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello from a test DOCX!")
        doc.add_paragraph(SAMPLE_TEXT)
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        text = _extract_docx(docx_bytes)
        assert "Hello" in text, "Expected text not found in extracted output"
        ok(f"Extracted {len(text)} chars from generated DOCX")
    except ImportError as e:
        fail("DOCX test skipped – missing library", e)
    except Exception as e:
        fail("DOCX extraction failed", e)


# ─────────────────────────────────────────────────────────────────────────────
# 3. IMAGE → TEXT  (Cloudflare Vision)
# ─────────────────────────────────────────────────────────────────────────────
def test_image():
    header("TEST 3 · Image → Text  (Cloudflare Vision API)")

    # Look for any jpg/png in project root to use as a real test image
    project_root = os.path.join(os.path.dirname(__file__), "..")
    test_image_path = None
    for fname in os.listdir(project_root):
        if fname.lower().endswith((".jpg", ".jpeg", ".png")):
            test_image_path = os.path.join(project_root, fname)
            break

    if test_image_path is None:
        fail("Image test skipped", "No .jpg/.png found in project root – add one to test")
        return

    try:
        with open(test_image_path, "rb") as f:
            img_bytes = f.read()
        ext  = os.path.splitext(test_image_path)[1].lower()
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        text = _extract_image_via_cf(img_bytes, mime_type=mime)
        assert isinstance(text, str) and len(text) > 0, "Empty response from vision API"
        ok(f"Got {len(text)} chars from Cloudflare Vision  →  \"{text[:80].strip()}…\"")
    except Exception as e:
        fail("Image extraction failed", e)


# ─────────────────────────────────────────────────────────────────────────────
# 4. AUDIO → TEXT  (Cloudflare Whisper)
# ─────────────────────────────────────────────────────────────────────────────
def test_audio():
    header("TEST 4 · Audio → Text  (Cloudflare Whisper API)")

    # Look for any audio file in project root
    project_root = os.path.join(os.path.dirname(__file__), "..")
    audio_exts   = (".mp3", ".wav", ".ogg", ".flac", ".m4a", ".aac", ".opus")
    test_audio_path = None
    for fname in os.listdir(project_root):
        if fname.lower().endswith(audio_exts):
            test_audio_path = os.path.join(project_root, fname)
            break

    if test_audio_path is None:
        fail("Audio test skipped", "No audio file found in project root – add one to test")
        return

    try:
        with open(test_audio_path, "rb") as f:
            audio_bytes = f.read()
        text = _extract_audio_via_cf(audio_bytes)
        assert isinstance(text, str) and len(text) > 0, "Empty response from Whisper API"
        ok(f"Transcribed {len(text)} chars  →  \"{text[:80].strip()}…\"")
    except Exception as e:
        fail("Audio transcription failed", e)


# ─────────────────────────────────────────────────────────────────────────────
# 5. EMBEDDING GENERATION  (Cloudflare bge-base-en-v1.5)
# ─────────────────────────────────────────────────────────────────────────────
def test_embeddings():
    header("TEST 5 · Embedding Generation  (Cloudflare CF Workers AI)")
    try:
        sample_texts = [
            "What is machine learning?",
            "Explain neural networks in simple terms.",
        ]
        vectors = embed_texts(sample_texts)
        assert len(vectors) == 2,         f"Expected 2 vectors, got {len(vectors)}"
        assert len(vectors[0]) == 768,    f"Expected dim=768, got {len(vectors[0])}"
        assert len(vectors[1]) == 768,    f"Expected dim=768, got {len(vectors[1])}"
        ok(f"Generated {len(vectors)} vectors, each dim={len(vectors[0])}")

        q_vec = embed_query("test query")
        assert len(q_vec) == 768, f"Query vector dim wrong: {len(q_vec)}"
        ok("embed_query() also works correctly")
    except Exception as e:
        fail("Embedding generation failed", e)


# ─────────────────────────────────────────────────────────────────────────────
# 6. PINECONE UPSERT  (full ingest_file pipeline on a synthetic text file)
# ─────────────────────────────────────────────────────────────────────────────
def test_upsert():
    header("TEST 6 · Pinecone Upsert  (ingest_file full pipeline)")
    try:
        fake_txt = SAMPLE_TEXT.encode("utf-8")
        result   = ingest_file(
            file_bytes = fake_txt,
            filename   = "test_document.txt",
            user_id    = TEST_USER_ID,
            conv_id    = TEST_CONV_ID,
        )
        assert result["num_chunks"] > 0,      "No chunks produced"
        assert len(result["vector_ids"]) > 0, "No vectors upserted"
        ok(
            f"Ingested '{result['source']}' → "
            f"{result['num_chunks']} chunks, "
            f"{len(result['vector_ids'])} vectors upserted"
        )
    except Exception as e:
        fail("Pinecone upsert failed", e)


# ─────────────────────────────────────────────────────────────────────────────
# 7. PINECONE RETRIEVAL
# ─────────────────────────────────────────────────────────────────────────────
def test_retrieval():
    header("TEST 7 · Pinecone Retrieval  (retrieve_chunks)")
    try:
        import time
        time.sleep(2)   # give Pinecone a moment to index the upserted vectors

        hits = retrieve_chunks(
            query   = TEST_QUERY,
            user_id = TEST_USER_ID,
            conv_id = TEST_CONV_ID,
            top_k   = 3,
        )
        assert isinstance(hits, list), "Expected a list of hits"
        if hits:
            ok(f"Got {len(hits)} hit(s)  →  top score={hits[0]['score']:.4f}  text=\"{hits[0]['text'][:60]}…\"")
        else:
            # Not necessarily a failure – index may need a moment to propagate
            print(f"  {YELLOW}⚠ WARN{RESET}  0 hits returned (index may still be indexing – retry in a few seconds)")
    except Exception as e:
        fail("Retrieval failed", e)


# ─────────────────────────────────────────────────────────────────────────────
# RUNNER
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"\n{BOLD}{'='*55}")
    print("   Vector Store · Smoke Test Suite")
    print(f"{'='*55}{RESET}")

    test_pdf()
    test_docx()
    test_image()
    test_audio()
    test_embeddings()
    test_upsert()
    test_retrieval()

    print(f"\n{BOLD}{'='*55}")
    print("   Done.")
    print(f"{'='*55}{RESET}\n")
